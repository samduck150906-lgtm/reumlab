# -*- coding: utf-8 -*-
"""
REUMLAB(름랩) 회사 소개서 · 포트폴리오 Word(.docx) 생성기.

콘텐츠 출처(단일 소스):
  - index.html            : 히어로·서비스·진행방식·인수인계·가격·FAQ 카피
  - script.js (PROJECTS)  : 익명화된 포트폴리오 16개 사례
  - content/templates.json: 사이트/연락처 기본 정보

실행:  python3 scripts/generate-portfolio-docx.py
출력:  REUMLAB_포트폴리오.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------------------------------------------------ 색상 팔레트
NAVY   = RGBColor(0x0E, 0x1A, 0x2B)   # 딥 네이비 (제목/강조)
INK    = RGBColor(0x22, 0x28, 0x33)   # 본문 잉크
ACCENT = RGBColor(0x2F, 0x5B, 0xEA)   # 시그니처 블루
ACC_DK = RGBColor(0x1E, 0x3A, 0x8A)   # 진한 블루
MUTED  = RGBColor(0x6B, 0x72, 0x80)   # 보조 텍스트 회색
LINE   = RGBColor(0xD8, 0xDE, 0xE8)   # 얇은 구분선
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GOOD   = RGBColor(0x15, 0x7F, 0x4B)   # 체크(초록)
BAD    = RGBColor(0xB4, 0x2A, 0x2A)   # 문제(적)

SHADE_HEAD = "0E1A2B"   # 섹션 헤더 배경
SHADE_SOFT = "F2F5FB"   # 옅은 배경
SHADE_BLUE = "E8EEFE"   # 블루 톤 배경
SHADE_ZEBRA = "F6F8FC"  # 표 얼룩

KO_FONT = "맑은 고딕"     # Malgun Gothic — 한글
EN_FONT = "맑은 고딕"

# ------------------------------------------------------------------ 저수준 헬퍼
def set_run_font(run, size=None, bold=None, color=None, font=KO_FONT, italic=None):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:eastAsia'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(el, hex_fill):
    """단락 또는 셀 요소에 배경색."""
    pr = el.get_or_add_tcPr() if el.tag.endswith('}tc') else el.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hex_fill)
    pr.append(sh)


def p_spacing(p, before=0, after=6, line=1.28):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_para(doc, text="", size=10.5, bold=False, color=INK, align=None,
             before=0, after=6, line=1.3, font=KO_FONT, italic=False):
    p = doc.add_paragraph()
    p_spacing(p, before, after, line)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, font=font, italic=italic)
    return p


def bottom_border(p, color="2F5BEA", size=6):
    pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), '4')
    b.set(qn('w:color'), color)
    pbdr.append(b)
    pr.append(pbdr)


def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def set_cell_bg(cell, hex_fill):
    shade(cell._tc, hex_fill)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def table_grid_borders(table, color="D8DEE8", size=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(size))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, size=9.5, bold=False, color=INK, align=None,
              before=1, after=1, font=KO_FONT, line=1.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p_spacing(p, before, after, line)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, font=font)
    return p


def cell_add_line(cell, text, size=9.5, bold=False, color=INK, align=None,
                  before=1, after=1, font=KO_FONT, line=1.2):
    p = cell.add_paragraph()
    p_spacing(p, before, after, line)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, font=font)
    return p


def section_header(doc, kicker, title, before=18):
    """진한 배경의 섹션 헤더 (kicker 라벨 + 큰 제목)."""
    kp = add_para(doc, kicker.upper(), size=8.5, bold=True, color=ACCENT,
                  before=before, after=1, line=1.0)
    kp.paragraph_format.keep_with_next = True
    tp = add_para(doc, title, size=17, bold=True, color=NAVY, before=0, after=4, line=1.1)
    tp.paragraph_format.keep_with_next = True
    bottom_border(tp, color="2F5BEA", size=8)
    add_para(doc, "", size=2, after=4)


def bullet(doc, text, label=None, size=10, color=INK, mark="•", mark_color=ACCENT,
           after=3, indent=0.2):
    p = doc.add_paragraph()
    p_spacing(p, 0, after, 1.28)
    p.paragraph_format.left_indent = Cm(indent + 0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    rm = p.add_run(mark + "  ")
    set_run_font(rm, size=size, bold=True, color=mark_color)
    if label:
        rl = p.add_run(label + " ")
        set_run_font(rl, size=size, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)
    return p


# ==================================================================== 데이터
SITE = {
    "brand": "REUMLAB",
    "brand_ko": "름랩",
    "studio": "앱 · 웹 · AI 개발 스튜디오",
    "ceo": "성아름",
    "biz_no": "793-12-03247",
    "addr": "경기 화성시 동탄첨단산업1로 58 (동탄첨단산업단지)",
    "tel": "010-8111-9370",
    "email": "ceo@eternalsix.com",
    "web": "reumlab.com",
}

INTRO = (
    "름랩(REUMLAB)은 아이디어를 실제로 운영 가능한 서비스로 만드는 앱·웹·AI 개발 스튜디오입니다. "
    "앱, SaaS, ERP, AI 자동화 기획부터 개발, 배포, 그리고 소스코드·운영 권한 인수인계까지 한 흐름으로 진행합니다. "
    "기술 용어가 아니라 화면과 사용자 흐름을 기준으로 소통하며, 비개발자 대표님이 만든 뒤에도 직접 운영할 수 있는 상태로 넘겨드리는 것을 원칙으로 합니다."
)

HERO_POINTS = [
    ("대표 직접 상담", "상담부터 구조 설계·최종 검수까지 대표가 직접 확인합니다."),
    ("단계별 진행 공유", "단계마다 확인 가능한 테스트 링크와 진행 내역을 공유합니다."),
    ("소스코드·권한 이관", "완료 후 소스코드와 도메인·배포·DB 등 주요 계정 권한을 대표님 소유로 이관합니다."),
]

# 목적별 서비스 (index.html SERVICES 섹션)
SERVICES = [
    {
        "name": "빠른 검증용 웹 MVP",
        "when": "아이디어를 빠르게 화면과 기능으로 구현해야 하는 경우",
        "feature": "랜딩·소개, 문의·예약, 간단 관리자",
        "scope": "화면 흐름 설계 · 반응형 · 기본 SEO · 소스코드 이관",
        "cases": "공간 예약·운영 웹 · 디지털 상품 마켓플레이스",
    },
    {
        "name": "모바일 앱 MVP",
        "when": "Flutter 기반 iOS·Android 앱이 필요한 경우",
        "feature": "회원·로그인, 목록·검색, 알림, 결제/예약",
        "scope": "앱 기획 · 크로스플랫폼 개발 · 관리자 연동 · 소스코드 이관",
        "cases": "지역 교육정보 탐색·리뷰 앱 · B2B 파트너·리드 관리 앱",
    },
    {
        "name": "운영관리 ERP · SaaS",
        "when": "회원·결제·예약·정산·상담 등 관리자 기능이 필요한 경우",
        "feature": "통합 관리자, 역할별 권한, 데이터 연동, 정산",
        "scope": "서비스 구조 설계 · 관리자 개발 · DB · 배포 · 이관",
        "cases": "교육기관 운영 ERP · 소상공인 운영 SaaS",
    },
    {
        "name": "AI 기능 및 업무 자동화",
        "when": "반복 작업·문서 생성·검색·분류·영업 업무를 자동화하려는 경우",
        "feature": "챗봇·요약·추천, 문서 자동 생성, 업무 자동화",
        "scope": "필요 기능 설계 · AI 연동 · 안정성 처리 · 이관",
        "cases": "AI 기반 영업 자동화 SaaS · AI 업무 자동화 허브",
    },
    {
        "name": "데이터 · SEO 시스템",
        "when": "대규모 데이터 수집·검색·인덱싱, 지역·서비스 페이지 자동화가 필요한 경우",
        "feature": "수집 파이프라인, 인덱싱·검색, 대량 페이지 자동화",
        "scope": "시스템 설계 · 빌드 자동화 · 구조화 데이터 · 이관",
        "cases": "대규모 데이터 수집·검색 시스템 · 대규모 SEO 페이지 자동화 시스템",
    },
]

# 익명 포트폴리오 (script.js PROJECTS)
PROJECTS = [
    {"title": "교육기관 운영 ERP", "chip": "ERP · 웹",
     "problem": "수납·출결·상담이 여러 채널에 흩어져, 매번 수기로 취합해야 했습니다.",
     "features": ["수강생·학부모 통합 관리", "수납·출결·상담 기록 일원화", "역할별 관리자 권한", "사용자 앱과 실시간 데이터 연동"],
     "scope": "서비스 구조 설계 · 관리자 웹 개발 · 데이터베이스 · 배포",
     "tech": "Flutter · 관리자 웹(React) · PostgreSQL(Supabase) · 권한(RLS) · 알림",
     "deliver": "전체 소스코드 · 관리자 웹+사용자 앱 · 실행·배포 문서 · 관리자 계정·권한 정리"},
    {"title": "지역 교육정보 탐색·리뷰 앱", "chip": "모바일 앱",
     "problem": "지역 교육 정보가 흩어져 있어, 조건에 맞는 곳을 비교하기 어려웠습니다.",
     "features": ["조건별 검색·필터", "지도 기반 탐색", "리뷰·평점", "맞춤 추천"],
     "scope": "앱 기획 · Flutter 앱 개발 · 추천 로직 · 관리자 연동",
     "tech": "Flutter · 검색·추천 로직 · 지도 SDK · PostgreSQL · 푸시 알림",
     "deliver": "전체 소스코드 · iOS·Android 앱 · 관리자 · 실행·배포 문서"},
    {"title": "생활서비스 매칭 플랫폼", "chip": "웹 · 앱",
     "problem": "견적 요청과 업체 배정이 전화·수기로 이뤄져 누락과 중복이 잦았습니다.",
     "features": ["요청서 기반 매칭", "업체·고객 양면 관리", "견적·정산 흐름", "문자·카카오 알림"],
     "scope": "플랫폼 구조 설계 · 웹/앱 개발 · 알림 연동 · 관리자",
     "tech": "Next.js · 모바일 앱 · PostgreSQL · 문자·카카오 알림",
     "deliver": "전체 소스코드 · 웹+앱+관리자 · 알림 연동 설정 · 실행·배포 문서"},
    {"title": "B2B 파트너·리드 관리 앱", "chip": "모바일 앱 · 관리자",
     "problem": "파트너별 리드 현황이 엑셀로만 관리돼 실시간 공유가 안 됐습니다.",
     "features": ["리드 등록·단계 관리", "파트너용 앱 + 본사 대시보드", "실적·정산 집계", "알림"],
     "scope": "모노레포 설계 · 앱 + 관리자 웹 · 공통 모듈 · 배포",
     "tech": "React Native/Expo · Next.js 관리자 · 공통(shared) 패키지 · PostgreSQL · 알림",
     "deliver": "전체 소스코드(모노레포) · 앱+관리자 · 실행·배포 문서 · 계정·권한 정리"},
    {"title": "AI 기반 영업 자동화 SaaS", "chip": "AI · SaaS",
     "problem": "반복되는 영업 리드 정리·메시지 작성에 시간이 과도하게 들었습니다.",
     "features": ["리드 관리·시퀀스 빌더", "AI 카피 초안 생성", "CRM·애널리틱스", "플랜별 요금·권한"],
     "scope": "서비스 기획 · 웹 개발 · AI 연동 · 결제·권한 설계",
     "tech": "Next.js(App Router) · TypeScript · AI 연동 · 결제 · 다국어(i18n)",
     "deliver": "전체 소스코드 · 관리자·요금 설정 · 실행·배포 문서 · 계정·키 정리"},
    {"title": "공간 예약·운영 웹", "chip": "웹 · 예약",
     "problem": "예약이 여러 채널로 들어와 이중 예약과 누락이 발생했습니다.",
     "features": ["실시간 예약 캘린더", "결제·환불 흐름", "운영자 관리자", "이용 안내 자동화"],
     "scope": "예약 구조 설계 · 웹 개발 · 결제 연동 · 관리자",
     "tech": "Next.js · TypeScript · 결제 연동 · PostgreSQL",
     "deliver": "전체 소스코드 · 예약 웹+관리자 · 결제 연동 설정 · 실행·배포 문서"},
    {"title": "소상공인 운영 SaaS", "chip": "SaaS",
     "problem": "예약·고객·매출 관리를 서로 다른 도구로 나눠 쓰느라 데이터가 흩어졌습니다.",
     "features": ["예약·고객·매출 통합", "간편 정산", "멀티 매장·권한", "모바일 대응"],
     "scope": "서비스 구조 설계 · 웹앱 개발 · 데이터 모델 · 배포",
     "tech": "Next.js · TypeScript · PostgreSQL · 반응형 웹앱",
     "deliver": "전체 소스코드 · 운영 웹앱+관리자 · 실행·배포 문서 · 계정·권한 정리"},
    {"title": "AI 업무 자동화 허브", "chip": "AI · 자동화",
     "problem": "반복 문서·요청 작성을 매번 사람이 직접 처리해 병목이 생겼습니다.",
     "features": ["업무별 자동 생성 템플릿", "입력 몇 개로 결과 산출", "결과 관리·재사용", "사용량 기반 과금"],
     "scope": "서비스 기획 · 웹 개발 · AI 연동 · 과금 설계",
     "tech": "Next.js · TypeScript · AI 연동 · 스키마 검증 · 레이트리밋",
     "deliver": "전체 소스코드 · 관리자·과금 설정 · 실행·배포 문서 · API 키 정리"},
    {"title": "SNS 콘텐츠 자동화 서비스", "chip": "AI · 자동화",
     "problem": "채널 콘텐츠를 매번 수작업으로 기획·제작하느라 발행이 밀렸습니다.",
     "features": ["주제 입력 → 초안 자동 생성", "이미지·문구 세트", "발행 캘린더", "톤 설정"],
     "scope": "서비스 기획 · 웹 개발 · AI 연동",
     "tech": "Next.js · TypeScript · AI 연동 · 이미지 처리",
     "deliver": "전체 소스코드 · 관리자 · 실행·배포 문서 · API 키 정리"},
    {"title": "공공정보 AI 탐색 서비스", "chip": "AI · 검색",
     "problem": "필요한 공공·지원 정보가 흩어져 있어, 조건에 맞는 항목을 찾기 어려웠습니다.",
     "features": ["대화형 조건 좁히기", "조건 매칭 검색", "결과 요약", "재시도·fallback 안정성"],
     "scope": "서비스 기획 · 웹 개발 · AI·검색 연동",
     "tech": "Next.js · TypeScript(strict) · AI 연동 · 검증(zod) · 레이트리밋",
     "deliver": "전체 소스코드 · 관리자 · 실행·배포 문서 · API 키 정리"},
    {"title": "디지털 상품 마켓플레이스", "chip": "웹 · 커머스",
     "problem": "디지털 상품 판매·정산을 수기로 처리해 운영 부담이 컸습니다.",
     "features": ["상품 등록·판매", "결제·정산", "구매자·판매자 관리", "셀프 편집 관리자"],
     "scope": "커머스 구조 설계 · 웹 개발 · 결제 연동 · 관리자",
     "tech": "Next.js · 결제 연동 · PostgreSQL · CMS(셀프 편집)",
     "deliver": "전체 소스코드 · 커머스 웹+관리자 · 결제 연동 설정 · 실행·배포 문서"},
    {"title": "견적·청구 문서 자동화 도구", "chip": "웹 · 자동화",
     "problem": "견적서·인보이스를 매번 수기로 작성해 실수와 재작업이 잦았습니다.",
     "features": ["항목 입력 → 문서 자동 생성", "템플릿·브랜딩", "PDF 출력·발송", "이력 관리"],
     "scope": "서비스 기획 · 웹 개발 · 문서 생성 로직",
     "tech": "Next.js · TypeScript · 문서 생성 로직 · PDF",
     "deliver": "전체 소스코드 · 관리자 · 실행·배포 문서"},
    {"title": "연구·문서 작성 보조 서비스", "chip": "AI · 문서",
     "problem": "자료 정리와 초안 작성에 반복 작업이 많아 시간이 오래 걸렸습니다.",
     "features": ["자료 기반 초안 보조", "구조·형식 정리", "인용·근거 관리", "결과 내보내기"],
     "scope": "서비스 기획 · 웹 개발 · AI 연동",
     "tech": "Next.js · TypeScript · AI 연동 · 문서 처리",
     "deliver": "전체 소스코드 · 관리자 · 실행·배포 문서"},
    {"title": "대규모 데이터 수집·검색 시스템", "chip": "데이터",
     "problem": "흩어진 대량 데이터를 수집·정리·검색할 방법이 없어 활용이 어려웠습니다.",
     "features": ["대규모 수집 파이프라인", "정제·인덱싱", "빠른 검색", "중복·오류 처리"],
     "scope": "시스템 설계 · 수집·인덱싱 개발 · 검색 API",
     "tech": "수집 파이프라인 · 인덱싱·검색 · 데이터베이스 · API",
     "deliver": "전체 소스코드 · 수집·검색 시스템 · 실행·배포 문서 · 운영 가이드"},
    {"title": "대규모 SEO 페이지 자동화 시스템", "chip": "데이터 · SEO",
     "problem": "지역·서비스별 페이지를 수작업으로 만들 수 없을 만큼 조합이 많았습니다.",
     "features": ["키워드 매트릭스 자동 생성", "대량 정적 페이지 빌드", "내부 링크 클러스터링", "구조화 데이터"],
     "scope": "시스템 설계 · 빌드 스크립트 · SEO 구조 · 자동 사이트맵",
     "tech": "Next.js 정적 export · Node 빌드 스크립트 · JSON-LD · 자동 sitemap/robots",
     "deliver": "전체 소스코드 · 빌드 자동화 · 실행·배포 문서 · 운영 가이드"},
]

# 진행 방식 (Process)
PROCESS = [
    ("01", "문의 및 가능 여부 검토",
     "받는 것 — 핵심 기능, 목표 일정, 참고 서비스, 예상 예산",
     "제공하는 것 — 구현 가능 여부, 권장 제작 범위, 예상 비용 구간",
     "산출물: 기능 정의 초안"),
    ("02", "범위 확정",
     "함께 정하는 것 — 포함 기능과 제외 범위를 나눠 확정",
     "제공하는 것 — 기능 목록, 화면 구조, 일정, 결제 단계",
     "산출물: 화면 구조도"),
    ("03", "화면 및 기능 제작",
     "진행 방식 — 코드부터 짜지 않고 화면 흐름부터 정리",
     "제공하는 것 — 확인 가능한 테스트 화면, 진행 내역, 주요 의사결정",
     "산출물: 테스트 링크"),
    ("04", "통합 테스트",
     "점검 범위 — 핵심 사용자 흐름, 모바일 반응형, 권한·오류",
     "제공하는 것 — QA 체크리스트와 수정 반영 내역",
     "산출물: QA 체크리스트"),
    ("05", "배포 및 인수인계",
     "넘겨드리는 것 — 운영 서비스, 소스코드, 주요 계정",
     "함께 드리는 것 — 운영 가이드, 수정 방법, 기본 장애 대응",
     "산출물: 인수인계 문서"),
]

HANDOVER = [
    "최종 소스코드", "배포 환경 및 프로젝트 권한", "데이터베이스 프로젝트 권한",
    "도메인 연결 정보", "관리자 계정", "주요 환경변수 목록",
    "운영 및 수정 가이드", "기본적인 장애 대응 방법", "계약 범위에 포함된 운영 교육",
]

WHY = [
    ("개발사를 바꾸고 싶어도 소스코드가 없다", "계약 범위의 소스코드 전체 이관 · 실행·배포 문서 포함"),
    ("어떤 계정으로 배포됐는지 알 수 없다", "도메인·배포·데이터베이스 계정 권한 정리 후 대표님 소유로 이관"),
    ("작은 문구 수정에도 추가 비용이 발생한다", "관리자에서 직접 수정 가능한 구조 · 콘텐츠 수정 AI 운영 교육 포함"),
    ("진행 상황을 결과물이 나올 때까지 알 수 없다", "단계별 테스트 링크와 진행 내용 공유 · 중간에 방향 조정"),
    ("기술 용어만 듣고 완성 모습을 이해하기 어렵다", "화면과 사용자 흐름 중심으로 설명 · 비개발자 기준 커뮤니케이션"),
    ("관리자 기능이 부족해 운영할 때마다 요청해야 한다", "운영에 필요한 관리자 기능과 운영 문서·AI 활용 가이드 제공"),
]

# 가격
WEB_PLANS = [
    ("웹 스타터", "소상공인 소개·문의용 랜딩", "980,000원", "약 5일",
     ["원페이지 랜딩 · 모바일 반응형", "문의·예약 CTA 연결", "소스코드 이관 · AI 수정 교육 1회"],
     "별도: 멀티페이지·복잡한 관리자·외부 연동"),
    ("웹 + 강력 마케팅  ★대표 추천", "검색·광고로 문의 늘리기", "1,960,000원", "약 10일",
     ["전환 카피 설계 + 기본 SEO", "GA·픽셀 세팅 + 광고 소재 1세트", "소스코드 이관 · AI 수정 교육 1회"],
     "별도: 유료 광고비 · 멀티페이지 확장"),
    ("웹 비즈니스", "멀티페이지·블로그 정식 웹", "3,800,000원", "약 14일",
     ["멀티페이지(5p 내외) + 블로그", "예약·문의 + 간단 CMS · 기본 SEO", "소스코드 이관 · AI 운영 1:1 교육"],
     "별도: 결제·회원 등 웹앱 기능"),
    ("웹 프리미엄", "관리자·외부 연동 웹앱", "5,800,000원", "약 21일",
     ["고도화 웹앱 + 관리자 페이지", "외부 연동(결제·지도·메일 1~2종)", "소스코드 이관 · AI 운영 교육"],
     "별도: 앱 · 대규모 연동 · AI 기능"),
]

APP_PLANS = [
    ("앱 라이트 MVP", "투자·테스트용 앱 빠른 검증", "5,800,000원", "약 14일",
     ["핵심 화면 3~5개 · 기본 데이터 연동", "문의·예약 흐름", "소스코드 이관 · AI 운영 1:1 교육"],
     "별도: 결제·회원 등 스탠다드 기능"),
    ("앱 스탠다드  ★가장 많이 선택", "회원·예약·결제 초기 서비스", "9,800,000원", "약 21일",
     ["회원/로그인 + DB + 결제 또는 예약", "기본 관리자 + 실행·수정 가이드", "소스코드 이관 · AI 운영 1:1 교육"],
     "별도: AI 기능 · 대규모 연동"),
    ("앱 AI", "AI로 운영 자동화하는 서비스", "13,800,000원", "약 30일",
     ["스탠다드 + AI 기능 1종(챗봇·추천·요약)", "업무 자동화 협의", "소스코드 이관 · AI 운영 교육"],
     "별도: AI 사용량(API) 실비 · 고도화"),
    ("앱 프리미엄", "다기능·AI 고도화·연동 다수", "19,800,000원", "약 45일",
     ["멀티 기능 + AI 고도화", "운영·관리자 흐름 + 외부 연동 다수", "소스코드 이관 · AI 운영 교육"],
     "별도: 대규모 트래픽·전용 인프라 실비"),
]

COMPARE = [
    ("가격 공개", "VAT 포함 정액 선공개", "견적 전까지 비공개"),
    ("소스코드", "전체 이관 · 대표님 소유", "미제공·부분 제공 잦음"),
    ("수정 비용", "관리자·직접 수정 교육으로 최소화", "수정마다 견적·과금"),
    ("계정 권한", "도메인·배포·DB 계정 이관", "개발사 명의로 남는 경우 잦음"),
    ("납품 후", "운영 가이드 + 1:1 교육", "문의 시 추가 비용"),
]

FAQ = [
    ("개발 지식이 없어도 의뢰할 수 있나요?",
     "가능합니다. 름랩은 비개발자 대표님을 주요 고객으로 삼습니다. 기술 용어 대신 화면과 사용자 흐름을 기준으로 설명드리고, 꼭 필요한 기능과 우선순위를 함께 정리해 드립니다."),
    ("아이디어만 있는 상태에서도 가능한가요?",
     "가능합니다. 완성된 기획서가 없어도 됩니다. 만들고 싶은 서비스와 꼭 필요한 기능, 참고 서비스만 있으면 상담에서 적절한 제작 범위를 함께 정리합니다."),
    ("소스코드는 모두 받을 수 있나요?",
     "네. 프로젝트 종료 시 계약 범위의 전체 소스코드와 실행·빌드·배포 방법 문서를 함께 전달합니다. 이후 다른 개발자에게 맡기거나 직접 운영할 수 있는 상태를 지향합니다."),
    ("도메인과 서버 계정은 누구 명의로 만들어지나요?",
     "가능하면 대표님(고객) 명의로 생성하거나, 대표님 소유로 이전하는 것을 원칙으로 합니다. 배포·데이터베이스·도메인 등 주요 계정 권한을 정리해 이관합니다."),
    ("디자인도 함께 진행하나요?",
     "네. 화면 흐름 설계와 UI 디자인을 개발과 함께 진행합니다. 보유하신 브랜드 자료(로고·컬러·폰트)가 있으면 반영 속도가 빨라집니다."),
    ("개발 중간에 진행 상황을 확인할 수 있나요?",
     "네. 단계별로 확인 가능한 테스트 화면과 진행 내역, 주요 의사결정 사항을 공유합니다. 완성 후 ‘생각한 것과 다르다’는 상황을 줄이기 위해 중간에 방향을 맞춥니다."),
    ("앱스토어 등록도 가능한가요?",
     "가능합니다. 다만 스토어 심사, 개발자 계정, 심사 기간, 정책 이슈는 플랫폼 사정에 따라 달라질 수 있어 일정과 비용은 상담 시 별도로 안내드립니다."),
    ("유지보수는 어떻게 진행되나요?",
     "기본 오류 확인 기간을 제공하며, 콘텐츠 수정은 관리자와 AI 운영 교육으로 직접 하실 수 있습니다. 이후 기능 추가나 운영 대행이 필요하면 별도 유지보수 또는 추가 개발로 협의합니다."),
    ("계약 후 추가 비용이 발생할 수 있나요?",
     "계약서에 확정된 개발 범위는 추가 비용 없이 진행합니다. 범위 밖의 신규 기능, 외부 서비스 실비(호스팅·결제·스토어 계정 등)는 진행 전 비용과 일정을 먼저 안내드립니다."),
    ("실제 고객사와 서비스명이 공개되지 않은 이유는 무엇인가요?",
     "고객사의 사업 정보와 서비스 전략을 보호하기 위해 일부 프로젝트명과 URL은 공개하지 않습니다. 대신 실제 구현 범위, 주요 기능, 화면과 개발 과정을 확인할 수 있도록 익명화된 사례로 제공합니다."),
]


# ==================================================================== 문서 조립
def build():
    doc = Document()

    # 기본 스타일
    normal = doc.styles['Normal']
    normal.font.name = KO_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), KO_FONT)
    rf.set(qn('w:ascii'), KO_FONT)
    rf.set(qn('w:hAnsi'), KO_FONT)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    page_w = sec.page_width - sec.left_margin - sec.right_margin

    # ------------------------------------------------ 표지
    add_para(doc, "", after=40)
    add_para(doc, "COMPANY PROFILE & PORTFOLIO", size=10, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    t = add_para(doc, "REUMLAB", size=52, bold=True, color=NAVY,
                 align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_para(doc, "름랩 · 앱 · 웹 · AI 개발 스튜디오", size=15, bold=True, color=INK,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    line_p = add_para(doc, "", after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    bottom_border(line_p, color="2F5BEA", size=10)

    add_para(doc, "아이디어를 실제로 운영 가능한 서비스로 만듭니다.", size=16, bold=True,
             color=ACC_DK, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "앱 · SaaS · ERP · AI 자동화 기획부터 배포와 소스코드 · 운영 권한 인수인계까지",
             size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)

    # 표지 하단 연락처 카드
    card = doc.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.columns[0].width = page_w
    c = card.cell(0, 0)
    c.width = page_w
    set_cell_bg(c, SHADE_HEAD)
    cell_margins(c, top=200, bottom=200, left=260, right=260)
    cell_text(c, "름랩 REUMLAB", size=13, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    cell_add_line(c, f"대표 {SITE['ceo']}   |   {SITE['tel']}   |   {SITE['email']}",
                  size=10.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    cell_add_line(c, f"{SITE['web']}   |   {SITE['addr']}",
                  size=9.5, color=RGBColor(0xC6, 0xD2, 0xE6),
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    doc.add_page_break()

    # ------------------------------------------------ 01 회사 소개
    section_header(doc, "About", "회사 소개", before=0)
    add_para(doc, INTRO, size=10.5, after=12, line=1.4)

    # 핵심 원칙 3개 카드 (3열 표)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(tbl)
    for i, (h, d) in enumerate(HERO_POINTS):
        cell = tbl.cell(0, i)
        cell.width = Emu(int(page_w / 3))
        set_cell_bg(cell, SHADE_BLUE)
        cell_margins(cell, top=150, bottom=150, left=150, right=150)
        cell_text(cell, h, size=11, bold=True, color=ACC_DK, after=3)
        cell_add_line(cell, d, size=9, color=INK, line=1.3)
    # 열 간격
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Emu(int(page_w / 3))
    add_para(doc, "", after=8)

    # 회사 정보 표
    info = [
        ("상호", "앱·웹·AI 개발 스튜디오 름랩 (REUMLAB)"),
        ("대표", SITE["ceo"]),
        ("사업자등록번호", SITE["biz_no"]),
        ("소재지", SITE["addr"]),
        ("연락처", f"{SITE['tel']}  ·  {SITE['email']}"),
        ("웹사이트", SITE["web"]),
        ("구축 경험", "앱·웹·SaaS·AI 15개 이상 구축 (공개 가능한 익명 사례 기준)"),
    ]
    it = doc.add_table(rows=len(info), cols=2)
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_grid_borders(it, color="E3E8F0", size=4)
    for r, (k, v) in enumerate(info):
        kc, vc = it.cell(r, 0), it.cell(r, 1)
        kc.width = Emu(int(page_w * 0.28)); vc.width = Emu(int(page_w * 0.72))
        set_cell_bg(kc, SHADE_SOFT)
        cell_margins(kc, 60, 60, 140, 100); cell_margins(vc, 60, 60, 140, 100)
        cell_text(kc, k, size=9.5, bold=True, color=NAVY)
        cell_text(vc, v, size=9.5, color=INK)

    # ------------------------------------------------ 02 서비스
    section_header(doc, "Services", "서비스 — 기술이 아니라 목적으로 나눴습니다")
    add_para(doc, "지금 필요한 목적에 가까운 것부터, 상담에서 범위를 함께 정합니다.",
             size=10, color=MUTED, after=10)
    for s in SERVICES:
        stbl = doc.add_table(rows=1, cols=1)
        stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = stbl.cell(0, 0); cell.width = page_w
        table_grid_borders(stbl, color="E3E8F0", size=4)
        cell_margins(cell, 120, 120, 160, 160)
        cell_text(cell, s["name"], size=12, bold=True, color=ACC_DK, after=2)
        cell_add_line(cell, s["when"], size=9.5, color=MUTED, after=5, line=1.3)
        for label, val in (("대표 기능", s["feature"]), ("제공 범위", s["scope"]), ("관련 사례", s["cases"])):
            p = cell.add_paragraph(); p_spacing(p, 0, 2, 1.3)
            rl = p.add_run(f"{label}   "); set_run_font(rl, size=9, bold=True, color=NAVY)
            rv = p.add_run(val); set_run_font(rv, size=9.5, color=INK)
        add_para(doc, "", after=4)

    # ------------------------------------------------ 03 포트폴리오
    doc.add_page_break()
    section_header(doc, "Portfolio", "포트폴리오 — 해결한 문제를 보여드립니다", before=0)
    add_para(doc,
             "실제 프로젝트명·고객사·서비스 URL·화면 속 개인정보는 고객사 요청에 따라 비공개 처리했습니다. "
             "대신 어떤 문제를 어떻게 설계하고 구현했는지 구체적으로 공개합니다.",
             size=9.5, color=MUTED, after=10, line=1.4)

    for idx, p in enumerate(PROJECTS, 1):
        ptbl = doc.add_table(rows=1, cols=1)
        ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = ptbl.cell(0, 0); cell.width = page_w
        table_grid_borders(ptbl, color="D8DEE8", size=4)
        cell_margins(cell, 130, 130, 170, 170)
        set_cell_bg(cell, SHADE_SOFT if idx % 2 else "FFFFFF")

        # 제목 줄: 번호 + 제목 + 칩
        hp = cell.paragraphs[0]; p_spacing(hp, 0, 3, 1.2)
        rn = hp.add_run(f"{idx:02d}  "); set_run_font(rn, size=12, bold=True, color=ACCENT)
        rt = hp.add_run(p["title"]); set_run_font(rt, size=12.5, bold=True, color=NAVY)
        rc = hp.add_run(f"     [ {p['chip']} ]"); set_run_font(rc, size=9, bold=True, color=ACC_DK)

        # 문제
        pp = cell.add_paragraph(); p_spacing(pp, 2, 4, 1.35)
        rpl = pp.add_run("문제   "); set_run_font(rpl, size=9, bold=True, color=BAD)
        rpv = pp.add_run(p["problem"]); set_run_font(rpv, size=9.5, color=INK)

        # 핵심 기능
        fp = cell.add_paragraph(); p_spacing(fp, 0, 4, 1.35)
        rfl = fp.add_run("핵심 기능   "); set_run_font(rfl, size=9, bold=True, color=GOOD)
        rfv = fp.add_run(" · ".join(p["features"])); set_run_font(rfv, size=9.5, color=INK)

        # 담당 범위 / 기술 / 산출물
        for label, val, col in (("담당 범위", p["scope"], NAVY),
                                ("기술 스택", p["tech"], NAVY),
                                ("산출물", p["deliver"], NAVY)):
            lp = cell.add_paragraph(); p_spacing(lp, 0, 2, 1.3)
            rl = lp.add_run(f"{label}   "); set_run_font(rl, size=8.5, bold=True, color=col)
            rv = lp.add_run(val); set_run_font(rv, size=9, color=RGBColor(0x3A,0x41,0x4D))
        add_para(doc, "", after=5)

    # ------------------------------------------------ 04 진행 방식
    doc.add_page_break()
    section_header(doc, "Process", "진행 방식 — 단계마다 무엇을 받는지 공개합니다", before=0)
    add_para(doc, "각 단계에서 대표님이 실제로 확인하게 되는 산출물을 함께 보여드립니다.",
             size=10, color=MUTED, after=10)
    for num, title, a, b, out in PROCESS:
        ptbl = doc.add_table(rows=1, cols=2)
        ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        no_table_borders(ptbl)
        numc, bodyc = ptbl.cell(0, 0), ptbl.cell(0, 1)
        numc.width = Emu(int(page_w * 0.13)); bodyc.width = Emu(int(page_w * 0.87))
        set_cell_bg(numc, SHADE_HEAD)
        cell_margins(numc, 120, 120, 60, 60); cell_margins(bodyc, 120, 120, 200, 120)
        cell_text(numc, num, size=18, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(bodyc, title, size=12, bold=True, color=NAVY, after=3)
        cell_add_line(bodyc, a, size=9.5, color=INK, after=2, line=1.3)
        cell_add_line(bodyc, b, size=9.5, color=INK, after=3, line=1.3)
        op = bodyc.add_paragraph(); p_spacing(op, 0, 0, 1.2)
        ro = op.add_run("▶ " + out); set_run_font(ro, size=9, bold=True, color=ACC_DK)
        add_para(doc, "", after=5)

    # 인수인계
    section_header(doc, "Handover", "인수인계 — 운영 권한까지 전달합니다")
    add_para(doc,
             "외주사에 종속되지 않도록, 결과물과 함께 아래 항목을 정리해 이관합니다. "
             "(프로젝트 범위와 사용하는 외부 서비스에 따라 세부 항목은 달라질 수 있습니다.)",
             size=9.5, color=MUTED, after=8, line=1.4)
    ht = doc.add_table(rows=3, cols=3)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(ht)
    for i, item in enumerate(HANDOVER):
        cell = ht.cell(i // 3, i % 3)
        cell.width = Emu(int(page_w / 3))
        set_cell_bg(cell, SHADE_BLUE)
        cell_margins(cell, 110, 110, 130, 90)
        pp = cell.paragraphs[0]; p_spacing(pp, 0, 0, 1.2)
        rk = pp.add_run("✓ "); set_run_font(rk, size=10, bold=True, color=GOOD)
        rv = pp.add_run(item); set_run_font(rv, size=9.5, bold=True, color=NAVY)
    for row in ht.rows:
        for cell in row.cells:
            cell.width = Emu(int(page_w / 3))

    # ------------------------------------------------ 05 우리가 다른 점
    doc.add_page_break()
    section_header(doc, "Why It Matters", "개발을 맡긴 뒤에도, 대표님이 서비스의 주인이어야 합니다", before=0)
    add_para(doc, "외주 개발에서 자주 겪는 문제와, 름랩이 그 문제를 해결하는 방식을 나란히 정리했습니다.",
             size=10, color=MUTED, after=10)
    wt = doc.add_table(rows=len(WHY) + 1, cols=2)
    wt.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_grid_borders(wt, color="E3E8F0", size=4)
    # 헤더
    h0, h1 = wt.cell(0, 0), wt.cell(0, 1)
    h0.width = Emu(int(page_w * 0.5)); h1.width = Emu(int(page_w * 0.5))
    set_cell_bg(h0, "F3E8E8"); set_cell_bg(h1, SHADE_BLUE)
    cell_margins(h0, 80, 80, 140, 100); cell_margins(h1, 80, 80, 140, 100)
    cell_text(h0, "✕  외주에서 자주 겪는 문제", size=9.5, bold=True, color=BAD)
    cell_text(h1, "✓  름랩의 해결 방식", size=9.5, bold=True, color=GOOD)
    for r, (prob, sol) in enumerate(WHY, 1):
        pc, sc = wt.cell(r, 0), wt.cell(r, 1)
        pc.width = Emu(int(page_w * 0.5)); sc.width = Emu(int(page_w * 0.5))
        cell_margins(pc, 80, 80, 140, 100); cell_margins(sc, 80, 80, 140, 100)
        cell_text(pc, prob, size=9.5, color=RGBColor(0x6b,0x4a,0x4a), line=1.3)
        cell_text(sc, sol, size=9.5, color=INK, line=1.3)

    # ------------------------------------------------ 06 가격
    doc.add_page_break()
    section_header(doc, "Pricing", "가격을 먼저 공개합니다", before=0)
    add_para(doc,
             "모든 금액은 VAT 포함 정액 기준입니다. 웹은 98만 원부터, 앱은 580만 원부터 — 필요한 깊이만큼 고르세요. "
             "패키지 범위를 넘는 기능은 상담 후 별도 견적으로 안내드립니다.",
             size=9.5, color=MUTED, after=12, line=1.4)

    def pricing_table(title, plans):
        add_para(doc, title, size=12, bold=True, color=NAVY, after=5)
        t = doc.add_table(rows=1, cols=4)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_grid_borders(t, color="D8DEE8", size=4)
        widths = [0.24, 0.18, 0.16, 0.42]
        heads = ["패키지", "금액(VAT 포함)", "기간", "포함 내용"]
        hdr = t.rows[0]
        for i, htext in enumerate(heads):
            cell = hdr.cells[i]; cell.width = Emu(int(page_w * widths[i]))
            set_cell_bg(cell, SHADE_HEAD)
            cell_margins(cell, 70, 70, 120, 90)
            cell_text(cell, htext, size=9, bold=True, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else None)
        for j, (name, sub, price, dur, incl, extra) in enumerate(plans):
            row = t.add_row()
            for i in range(4):
                row.cells[i].width = Emu(int(page_w * widths[i]))
                cell_margins(row.cells[i], 80, 80, 120, 90)
                if j % 2 == 0:
                    set_cell_bg(row.cells[i], SHADE_ZEBRA)
            cell_text(row.cells[0], name, size=9.5, bold=True, color=ACC_DK, after=1)
            cell_add_line(row.cells[0], sub, size=8, color=MUTED, line=1.2)
            cell_text(row.cells[1], price, size=10, bold=True, color=NAVY,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_text(row.cells[2], dur, size=9, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
            first = True
            for line in incl:
                if first:
                    cell_text(row.cells[3], "· " + line, size=8.8, color=INK, line=1.25, after=1)
                    first = False
                else:
                    cell_add_line(row.cells[3], "· " + line, size=8.8, color=INK, line=1.25, after=1)
            cell_add_line(row.cells[3], extra, size=8, color=MUTED, line=1.2)
        add_para(doc, "", after=8)

    pricing_table("웹(Web) 라인 — 가볍게 시작해 필요한 만큼 확장", WEB_PLANS)
    pricing_table("앱(App) 라인 — Flutter로 iOS·안드로이드 동시, 기능 깊이로 선택", APP_PLANS)
    add_para(doc,
             "※ 계약서에 확정된 개발 범위는 추가 비용 없이 진행합니다. 범위 밖의 신규 기능은 진행 전 비용과 일정을 먼저 안내하며, "
             "호스팅·결제·스토어 계정 등 외부 서비스 실비는 사전 고지 후 대표님 명의로 진행합니다.",
             size=8.5, color=MUTED, after=12, line=1.4)

    # 정액제 비교
    add_para(doc, "정액제가 남기는 차이", size=12, bold=True, color=NAVY, after=5)
    ct = doc.add_table(rows=1, cols=3)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_grid_borders(ct, color="D8DEE8", size=4)
    cw = [0.24, 0.42, 0.34]
    for i, htext in enumerate(["항목", "름랩", "일반 외주"]):
        cell = ct.rows[0].cells[i]; cell.width = Emu(int(page_w * cw[i]))
        set_cell_bg(cell, SHADE_HEAD); cell_margins(cell, 70, 70, 120, 90)
        cell_text(cell, htext, size=9, bold=True, color=WHITE,
                  align=None if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for j, (k, good, bad) in enumerate(COMPARE):
        row = ct.add_row()
        for i in range(3):
            row.cells[i].width = Emu(int(page_w * cw[i]))
            cell_margins(row.cells[i], 70, 70, 120, 90)
        set_cell_bg(row.cells[0], SHADE_SOFT)
        if j % 2 == 0:
            set_cell_bg(row.cells[1], SHADE_ZEBRA); set_cell_bg(row.cells[2], SHADE_ZEBRA)
        cell_text(row.cells[0], k, size=9.5, bold=True, color=NAVY)
        gp = row.cells[1].paragraphs[0]; p_spacing(gp, 1, 1, 1.25)
        rg = gp.add_run("✓ "); set_run_font(rg, size=9.5, bold=True, color=GOOD)
        rgt = gp.add_run(good); set_run_font(rgt, size=9.3, color=INK)
        cell_text(row.cells[2], bad, size=9.3, color=MUTED)

    # ------------------------------------------------ 07 대표 소개
    doc.add_page_break()
    section_header(doc, "Who We Are", "누가 만드는지, 먼저 밝히고 시작합니다", before=0)
    ceo_tbl = doc.add_table(rows=1, cols=1)
    ceo_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = ceo_tbl.cell(0, 0); cc.width = page_w
    set_cell_bg(cc, SHADE_SOFT)
    table_grid_borders(ceo_tbl, color="E3E8F0", size=4)
    cell_margins(cc, 160, 160, 200, 200)
    cell_text(cc, f"{SITE['ceo']}  ·  름랩 대표", size=13, bold=True, color=NAVY, after=2)
    cell_add_line(cc, "경기 화성 동탄첨단산업단지 · 앱·웹·AI MVP 개발 스튜디오",
                  size=9, color=MUTED, after=8)
    cell_add_line(cc,
                  "외주로 서비스를 만든 뒤, 문구 하나 바꾸는 데도 매번 개발사에 연락하고 비용·시간을 쓰는 "
                  "대표님들을 많이 봤습니다. 름랩을 시작한 이유는 단순합니다. 만들어 드린 결과물을 대표님이 "
                  "직접 운영할 수 있는 상태로 넘겨드리기 위해서입니다.",
                  size=10, color=INK, after=6, line=1.45)
    cell_add_line(cc,
                  "그래서 상담부터 구조 설계와 최종 검수까지 대표가 직접 확인하고, 완료 후에는 소스코드와 운영 "
                  "권한을 이관하며, AI로 간단한 수정을 직접 하실 수 있게 1:1로 교육합니다.",
                  size=10, color=INK, after=8, line=1.45)
    for badge in (f"대표 실명 공개 · 사업자등록 {SITE['biz_no']}",
                  f"사업장 주소 공개 · {SITE['addr']}",
                  "상담부터 구조 설계·최종 검수까지 대표가 직접 확인"):
        bp = cc.add_paragraph(); p_spacing(bp, 0, 2, 1.3)
        rb = bp.add_run("• "); set_run_font(rb, size=9.5, bold=True, color=ACCENT)
        rbt = bp.add_run(badge); set_run_font(rbt, size=9.3, color=INK)

    # ------------------------------------------------ 08 FAQ
    section_header(doc, "FAQ", "외주개발 전에 자주 묻는 질문")
    for i, (q, a) in enumerate(FAQ, 1):
        qp = add_para(doc, "", after=2)
        rq = qp.add_run(f"Q{i}. "); set_run_font(rq, size=10.5, bold=True, color=ACCENT)
        rqt = qp.add_run(q); set_run_font(rqt, size=10.5, bold=True, color=NAVY)
        qp.paragraph_format.keep_with_next = True
        ap = add_para(doc, "", after=9, line=1.4)
        ap.paragraph_format.left_indent = Cm(0.55)
        ra = ap.add_run(a); set_run_font(ra, size=9.8, color=RGBColor(0x3A,0x41,0x4D))

    # ------------------------------------------------ 09 문의(CTA)
    doc.add_page_break()
    add_para(doc, "", after=30)
    add_para(doc, "CONTACT", size=10, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_para(doc, "만들고 싶은 서비스를 알려주세요.", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "완성된 기획서가 없어도 괜찮습니다. 필요한 기능과 목표를 적어주시면 가능한 범위부터 먼저 정리해 드립니다.",
             size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    for line in ("문의만으로 계약이 진행되지 않습니다",
                 "범위와 비용을 확인한 뒤 결정할 수 있습니다",
                 "전달한 내용은 견적 검토 목적으로만 사용합니다"):
        cp = add_para(doc, "", after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
        rc = cp.add_run("✓  " + line); set_run_font(rc, size=10, color=INK)
    add_para(doc, "", after=20)

    contact = doc.add_table(rows=1, cols=1)
    contact.alignment = WD_TABLE_ALIGNMENT.CENTER
    contact.columns[0].width = page_w
    cc2 = contact.cell(0, 0); cc2.width = page_w
    set_cell_bg(cc2, SHADE_HEAD)
    cell_margins(cc2, 220, 220, 260, 260)
    cell_text(cc2, "름랩 REUMLAB", size=15, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    cell_add_line(cc2, f"전화 · 카카오   {SITE['tel']}", size=12, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    cell_add_line(cc2, f"이메일   {SITE['email']}", size=11, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    cell_add_line(cc2, f"웹사이트   {SITE['web']}", size=11, color=RGBColor(0xC6,0xD2,0xE6),
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    add_para(doc, "", after=24)
    add_para(doc, f"대표 {SITE['ceo']}  ·  사업자등록 {SITE['biz_no']}  ·  {SITE['addr']}",
             size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "© 2026 이터널식스 (REUMLAB). All Rights Reserved.",
             size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    out_path = "REUMLAB_포트폴리오.docx"
    doc.save(out_path)
    print("saved:", out_path)


if __name__ == "__main__":
    build()
